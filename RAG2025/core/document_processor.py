#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
文档处理模块 - 处理文档分割和处理相关功能
"""

import os
import uuid
import tempfile
import io
from typing import List, Dict, Any, Optional, Union
from haystack import Document

from langchain.text_splitter import RecursiveCharacterTextSplitter as LangChainSplitter

from .config import logger

def split_text_into_fragments(text, fragment_size=20):
    """将文本分割成固定大小的片段"""
    words = text.split()
    fragments = []
    for i in range(0, len(words), fragment_size):
        fragment = ' '.join(words[i:i+fragment_size])
        fragments.append(fragment)
    return fragments

def process_file(file):
    """处理上传的文件，提取文本内容"""
    filename = file.filename
    file_extension = os.path.splitext(filename)[1].lower()
    
    # 创建临时文件
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp:
        file.save(temp.name)
        temp_path = temp.name
    
    try:
        # 根据文件类型处理
        if file_extension in ['.txt', '.md']:
            # 文本文件直接读取
            with open(temp_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            # 创建Document对象
            doc = Document(
                id=str(uuid.uuid4()),
                content=content,
                meta={
                    "filename": filename,
                    "type": file_extension[1:].upper(),
                    "source": "用户上传"
                }
            )
            return {"success": True, "document": doc}
            
        elif file_extension in ['.pdf']:
            # 处理PDF文件
            try:
                import fitz  # PyMuPDF
                
                text_parts = []
                with fitz.open(temp_path) as pdf:
                    for page_num in range(len(pdf)):
                        page = pdf[page_num]
                        text_parts.append(page.get_text())
                
                content = "\n\n".join(text_parts)
                
                # 创建Document对象
                doc = Document(
                    id=str(uuid.uuid4()),
                    content=content,
                    meta={
                        "filename": filename,
                        "type": "PDF",
                        "pages": len(text_parts),
                        "source": "用户上传"
                    }
                )
                return {"success": True, "document": doc}
            except ImportError:
                return {"success": False, "error": "处理PDF文件需要安装PyMuPDF库"}
            except Exception as e:
                return {"success": False, "error": f"处理PDF文件时出错: {str(e)}"}
                
        elif file_extension in ['.doc', '.docx']:
            # 处理Word文件
            try:
                import docx
                
                doc_obj = docx.Document(temp_path)
                content = "\n\n".join([para.text for para in doc_obj.paragraphs if para.text])
                
                # 创建Document对象
                doc = Document(
                    id=str(uuid.uuid4()),
                    content=content,
                    meta={
                        "filename": filename,
                        "type": file_extension[1:].upper(),
                        "source": "用户上传"
                    }
                )
                return {"success": True, "document": doc}
            except ImportError:
                return {"success": False, "error": "处理Word文件需要安装python-docx库"}
            except Exception as e:
                return {"success": False, "error": f"处理Word文件时出错: {str(e)}"}
        else:
            return {"success": False, "error": f"不支持的文件类型: {file_extension}"}
    finally:
        # 清理临时文件
        try:
            os.unlink(temp_path)
        except:
            pass

class TextSplitter:
    """文本分割器，将长文本分割成较小的块"""
    
    def __init__(self, chunk_size=1000, chunk_overlap=200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.splitter = LangChainSplitter(
            chunk_size=self.chunk_size,
            chunk_overlap=self.chunk_overlap,
            separators=["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""]
        )
    
    def split_document(self, document):
        """将文档分割成较小的块"""
        if not document or not document.content:
            return []
            
        # 使用LangChain分割器分割文本
        text_chunks = self.splitter.split_text(document.content)
        
        # 创建新的Document对象
        documents = []
        
        # 首先，将原始文档标记为父文档
        parent_meta = document.meta.copy() if document.meta else {}
        parent_meta["is_fragment"] = False  # 明确标记为非分段
        parent_meta["has_fragments"] = True  # 标记为有分段
        parent_meta["total_fragments"] = len(text_chunks)  # 记录总分段数
        
        # 更新原始文档的元数据
        document.meta = parent_meta
        
        # 添加父文档到结果列表
        documents.append(document)
        
        # 然后创建分段文档
        for i, chunk in enumerate(text_chunks):
            # 创建新的ID，但保持与原始文档的关联
            chunk_id = f"{document.id}_chunk_{i}"
            
            # 复制元数据并添加分块信息
            meta = document.meta.copy() if document.meta else {}
            meta["chunk_id"] = i
            meta["total_chunks"] = len(text_chunks)
            meta["parent_id"] = document.id
            meta["is_fragment"] = True  # 明确标记为分段
            
            # 创建新的Document对象
            doc = Document(
                id=chunk_id,
                content=chunk,
                meta=meta
            )
            documents.append(doc)
        
        return documents 